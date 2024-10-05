# -*- encoding: utf-8 -*-

# Import necessary modules from Django and other libraries
import os
import json

from django import template
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse, HttpResponseRedirect, FileResponse
from django.shortcuts import render
from django.template import loader
from django.urls import reverse
from docx import settings
from django.utils import timezone

from django.shortcuts import render, get_object_or_404, redirect
from .models import WendlerPlan
from .forms import WendlerPlanForm

# Import PDF generation libraries
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import Table, TableStyle, Paragraph
from reportlab.pdfgen.canvas import Canvas
from reportlab.lib.styles import getSampleStyleSheet

# Import Word document generation library
from docx import Document
from docx.shared import Inches

# Import custom form
from .forms import WendlerForm
from .models import WendlerPlan
from reportlab.pdfgen import canvas
import io
from django.conf import settings

from django.shortcuts import render, get_object_or_404, redirect
from .models import WendlerPlan
from .forms import WendlerPlanForm

# View for the dashboard page, requires login
@login_required(login_url="/login/")
def index(request):
    context = {'segment': 'index'}
    html_template = loader.get_template('home/dashboard.html')
    return HttpResponse(html_template.render(context, request))


import csv
from django.shortcuts import render
import plotly.graph_objs as go
from plotly.offline import plot


# def plot_csv(request):


# View for the Wendler 5/3/1 calculator
def wendler_view(request):
    if request.method == 'POST':
        form = WendlerForm(request.POST)
        if form.is_valid():
            # Get the one rep max weight from the form
            number = int(request.POST['weight'])
            global global_wendler_list

            # List of percentages for the Wendler 5/3/1 program
            percentage_list = [0.40, 0.50, 0.60, 0.65, 0.70, 0.75, 0.80, 0.85, 0.90, 0.95]

            # Calculate weights for each percentage
            calculated_dict = {}
            for num in percentage_list:
                calc_num = (number * num - 20) / 2
                # Round to nearest 2.5 kg
                if (calc_num % 2.5) < 1.25:
                    calc_num = calc_num - (calc_num % 2.5)
                else:
                    calc_num = calc_num + 2.5 - (calc_num % 2.5)
                # Replace negative values with zero
                if calc_num < 0:
                    calc_num = 0
                calculated_dict[num] = calc_num

            # Create a dictionary with the exercise plan for each week
            exercise_dict = {
                'Week 1': {'Set 1': str(calculated_dict[0.4]) + 'kgx5',
                           'Set 2': str(calculated_dict[0.65]) + 'kgx5',
                           'Set 3': str(calculated_dict[0.75]) + 'kgx5',
                           'Set 4': str(calculated_dict[0.85]) + 'kgx5'},

                'Week 2': {'Set 1': str(calculated_dict[0.4]) + 'kgx3',
                           'Set 2': str(calculated_dict[0.7]) + 'kgx3',
                           'Set 3': str(calculated_dict[0.8]) + 'kgx3',
                           'Set 4': str(calculated_dict[0.9]) + 'kgx3'},

                'Week 3': {'Set 1': str(calculated_dict[0.4]) + 'kgx5',
                           'Set 2': str(calculated_dict[0.75]) + 'kgx5',
                           'Set 3': str(calculated_dict[0.85]) + 'kgx3',
                           'Set 4': str(calculated_dict[0.95]) + 'kgx1'},

                'Week 4': {'Set 1': str(calculated_dict[0.4]) + 'kgx5',
                           'Set 2': str(calculated_dict[0.4]) + 'kgx5',
                           'Set 3': str(calculated_dict[0.5]) + 'kgx5',
                           'Set 4': str(calculated_dict[0.6]) + 'kgx5'},
            }

            global_wendler_list = exercise_dict

            # Save the plan to the database
            plan = WendlerPlan.objects.create(
                user=request.user,
                name=form.cleaned_data.get('name', 'Default Wendler Plan'),
                weight=number
            )
            # You can create a related model to store `exercise_dict` if needed
            # For simplicity, let's assume you serialize it as JSON
            import json
            plan.exercise_data = json.dumps(exercise_dict)
            plan.save()

            return render(request, 'home/wendler.html',
                          {'form': form, 'number': number, 'calculated_dict': exercise_dict})
    else:
        form = WendlerForm()

    return render(request, 'home/wendler.html', {'form': form})

# View to generate PDF of the Wendler plan
def some_view(request):
    # Set up PDF styles and buffer
    styles = getSampleStyleSheet()
    style = styles["BodyText"]
    buffer = io.BytesIO()
    canv = canvas.Canvas(buffer)

    # Prepare data for the PDF table
    empty_list = []
    empty_list.append(['Week No.', 'Set 1', "Set 2", "Set 3", "Set 4"])
    # Add data for each week
    for week in range(1, 5):
        empty_list.append([f'Week {week}'] + list(global_wendler_list[f'Week {week}'].values()))

    # Create header and table for PDF
    header = Paragraph("<bold><font size=18>Wendler Exercise List</font></bold>", style)
    data = empty_list
    t = Table(data)
    # Set table styles
    t.setStyle(TableStyle([("BOX", (0, 0), (-1, -1), 0.25, colors.black),
                           ('INNERGRID', (0, 0), (-1, -1), 0.25, colors.black)]))
    # Add alternating row colors
    for each in range(len(data)):
        bg_color = colors.whitesmoke if each % 2 == 0 else colors.lightgrey
        t.setStyle(TableStyle([('BACKGROUND', (0, each), (-1, each), bg_color)]))

    # Draw header and table on the PDF
    aW, aH = 540, 720
    w, h = header.wrap(aW, aH)
    header.drawOn(canv, 72, 800)
    aH = aH - h
    w, h = t.wrap(aW, aH)
    t.drawOn(canv, 72, aH - h)
    canv.save()

    # Return the PDF as a file response
    buffer.seek(0)
    return FileResponse(buffer, as_attachment=True, filename='WendlerSheet.pdf')

def wendler_plan_list(request):
    wendler_plans = WendlerPlan.objects.filter(user=request.user)
    for plan in wendler_plans:
        if isinstance(plan.exercise_data, str):  # Check if it's a string
            try:
                plan.exercise_data = json.loads(plan.exercise_data)  # Parse JSON string
            except json.JSONDecodeError:
                plan.exercise_data = {}  # Set to empty dict if JSON is invalid
    return render(request, 'home/settings.html', {'wendler_plans': wendler_plans})



def update_wendler_plan(request, plan_id):
    plan = get_object_or_404(WendlerPlan, id=plan_id, user=request.user)

    if request.method == 'POST':
        form = WendlerPlanForm(request.POST, instance=plan)
        if form.is_valid():
            updated_plan = form.save(commit=False)  # Save the form but don't commit yet
            updated_plan.updated_at = timezone.now()  # Set updated_at to current time
            updated_plan.save()  # Now save the instance
            return redirect('wendler_plan_list')  # Redirect to the list of Wendler plans
    else:
        form = WendlerPlanForm(instance=plan)  # Pre-fill the form with the current data

    return render(request, 'home/update_wendler_plan.html', {'form': form, 'plan': plan})

def delete_wendler_plan(request, plan_id):
    plan = get_object_or_404(WendlerPlan, id=plan_id, user=request.user)  # Ensure the user owns the plan
    if request.method == 'POST':
        plan.delete()  # Delete the WendlerPlan object
        return redirect('wendler_plan_list')  # Redirect to the list view
    return render(request, 'home/confirm_delete.html', {'plan': plan})

# View to generate Word document of the Wendler plan
def word_doc_view(request):
    document = Document()
    docx_title = "WendlerSheet.docx"
    # Add content to the Word document
    document.add_paragraph("Wendler Exercise List")
    for week in range(1, 5):
        document.add_paragraph(f'Week {week}' + str(global_wendler_list[f'Week {week}'])[1:-1])
    document.add_page_break()

    # Prepare document for download
    f = io.BytesIO()
    document.save(f)
    length = f.tell()
    f.seek(0)
    response = HttpResponse(
        f.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename=' + docx_title
    response['Content-Length'] = length
    return response

# View for handling other pages, requires login
@login_required(login_url="/login/")
def pages(request):
    # List of CSV files
    csv_files = [
        'benchpress.csv',
        'age_skill_levels_deadlift.csv',
        'age_skill_levels_overheadpress.csv',
        'age_skill_levels_squat.csv'
    ]

    plot_divs = []  # This will hold all the plotly divs
    file_errors = []  # To keep track of any missing files

    # Prepare data for each CSV file
    for file_name in csv_files:
        file_path = os.path.join(settings.DATA_DIR, 'apps', 'dataset', file_name)
        try:
            with open(file_path, 'r') as file:
                csv_reader = csv.DictReader(file)
                csv_data = list(csv_reader)  # Read the data into a list

                # Prepare data for Plotly
                age = [int(row['Age']) for row in csv_data]  # Extract ages
                categories = ['Beginner', 'Novice', 'Intermediate', 'Advanced', 'Elite']

                traces = []
                for category in categories:
                    y = [int(row[category]) for row in csv_data if category in row]
                    trace = go.Bar(
                        x=age,
                        y=y,
                        name=f'{file_name[:-4]} - {category}',  # Name includes the type of lift
                        text=[f'{category} ({row["Age"]})' for row in csv_data],
                        hoverinfo='x+y+name',  # Shows Age, Weight, and Category on hover
                        hovertemplate='<b>Age:</b> %{x}<br><b>Weight:</b> %{y} kg<br><b>Category:</b> %{name}<extra></extra>',
                    )
                    traces.append(trace)

                # Create the figure for this specific lift
                layout = go.Layout(
                    title=f'Performance by Age and Skill Level for {file_name[:-4]}',
                    xaxis=dict(title='Age'),
                    yaxis=dict(title='Weight (kg)'),
                    barmode='group',  # Group bars for comparison
                )
                fig = go.Figure(data=traces, layout=layout)

                # Convert the figure to HTML and add to plot_divs
                plot_div = plot(fig, output_type='div', include_plotlyjs=True)
                plot_divs.append(plot_div)

        except FileNotFoundError:
            file_errors.append(f'Data file not found: {file_name}')

    # Context with multiple plot divs
    context = {
        'plot_divs': plot_divs,
        'file_errors': file_errors if file_errors else None
    }

    return render(request, 'home/transactions.html', context=context)

    context = {}
    try:
        load_template = request.path.split('/')[-1]
        if load_template == 'admin':
            return HttpResponseRedirect(reverse('admin:index'))
        context['segment'] = load_template
        html_template = loader.get_template('home/' + load_template)
        return HttpResponse(html_template.render(context, request))
    except template.TemplateDoesNotExist:
        # Handle 404 error
        html_template = loader.get_template('home/page-404.html')
        return HttpResponse(html_template.render(context, request))
    except:
        # Handle 500 error
        html_template = loader.get_template('home/page-500.html')
        return HttpResponse(html_template.render(context, request))